"""DOCX text extraction — mirrors pdf_parser.py's contract (accepts a
file-like object, returns one newline-joined string) so api/resume.py
can pick either extractor based on the uploaded file's real type
without the rest of the ingestion pipeline caring which one ran.

Requires the `python-docx` package (import name: `docx`). Add it to
requirements.txt / pyproject.toml — not modified here since that file
wasn't provided.
"""
import docx


def extract_text_from_docx(file_bytes) -> str:
    """`file_bytes` is a file-like object (e.g. io.BytesIO(raw_bytes)).
    Walks paragraphs AND tables — resume templates frequently use tables
    for two-column layouts or skill grids, and a paragraph-only walk
    silently drops that content entirely.
    """
    document = docx.Document(file_bytes)
    parts: list[str] = []

    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)